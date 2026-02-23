#!/usr/bin/env python3
"""
sync_summary.py - Automate SESSION_SUMMARY.md conversion to .docx

This script converts the SESSION_SUMMARY.md file to a Word document (.docx)
for easy sharing via WhatsApp or other platforms.

Usage:
    python sync_summary.py [--output OUTPUT_PATH]

Requirements:
    pip install python-docx

According to kuro-rules:
    "The user manually copies SESSION_SUMMARY.md to a Word document and WhatsApp.
     When creating a session summary, you MUST also generate or update a script
     that automates converting the markdown to .docx"
"""

import argparse
import re
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("ERROR: python-docx not installed.")
    print("Install with: pip install python-docx")
    sys.exit(1)


def parse_session_summary(markdown_content: str) -> list[dict]:
    """
    Parse SESSION_SUMMARY.md into structured session data.
    
    Args:
        markdown_content: Raw markdown content of SESSION_SUMMARY.md
        
    Returns:
        List of session dictionaries with parsed data
    """
    sessions = []
    
    # Split by session headers
    session_pattern = r'# Session Summary — (\d{4}-\d{2}-\d{2})(?:\s+\(Part\s+\d+\))?'
    parts = re.split(session_pattern, markdown_content)
    
    # parts[0] is empty or content before first session
    # Then alternating: date, content, date, content...
    for i in range(1, len(parts), 2):
        if i + 1 < len(parts):
            date = parts[i]
            content = parts[i + 1]
            
            session = {
                'date': date,
                'editor': '',
                'francais': {},
                'english': {},
                'tests': '',
                'blockers': '',
                'progress': ''
            }
            
            # Extract editor
            editor_match = re.search(r'\*\*Editor\*\*:\s*(.+)', content)
            if editor_match:
                session['editor'] = editor_match.group(1).strip()
            
            # Extract French section
            fr_match = re.search(r'## Francais\s*(.*?)(?=## English|## ---|\*\*Tests\*\*|$)', content, re.DOTALL)
            if fr_match:
                session['francais'] = parse_section(fr_match.group(1))
            
            # Extract English section
            en_match = re.search(r'## English\s*(.*?)(?=\*\*Tests\*\*|$)', content, re.DOTALL)
            if en_match:
                session['english'] = parse_section(en_match.group(1))
            
            # Extract tests, blockers, progress
            tests_match = re.search(r'\*\*Tests\*\*:\s*(.+)', content)
            if tests_match:
                session['tests'] = tests_match.group(1).strip()
            
            blockers_match = re.search(r'\*\*Blockers\*\*:\s*(.+)', content)
            if blockers_match:
                session['blockers'] = blockers_match.group(1).strip()
            
            progress_match = re.search(r'\*\*Progress\*\*:\s*(\d+%)', content)
            if progress_match:
                session['progress'] = progress_match.group(1)
            
            sessions.append(session)
    
    return sessions


def parse_section(section_content: str) -> dict:
    """Parse a language section (Francais or English) into structured data."""
    result = {
        'what_done': [],
        'initiatives': [],
        'files_changed': [],
        'next_steps': []
    }
    
    # Map patterns to keys
    patterns = {
        'what_done': r'\*\*(?:Ce qui a ete fait|What was done)\*\*\s*:\s*(.*?)(?=\*\*Initiatives|\*\*Fichiers|\*\*Files|\*\*Etapes|\*\*Next|$)',
        'initiatives': r'\*\*(?:Initiatives donnees|Initiatives given)\*\*\s*:\s*(.*?)(?=\*\*Fichiers|\*\*Files|\*\*Etapes|\*\*Next|$)',
        'files_changed': r'\*\*(?:Fichiers modifies|Files changed)\*\*\s*:\s*(.*?)(?=\*\*Etapes|\*\*Next|$)',
        'next_steps': r'\*\*(?:Etapes suivantes|Next steps)\*\*\s*:\s*(.*?)(?=\*\*Tests|\*\*Blockers|\*\*Progress|$)'
    }
    
    for key, pattern in patterns.items():
        match = re.search(pattern, section_content, re.DOTALL)
        if match:
            items = match.group(1)
            # Extract bullet points
            bullets = re.findall(r'-\s*(.+)', items)
            result[key] = [b.strip() for b in bullets if b.strip()]
    
    return result


def create_document(sessions: list[dict]) -> Document:
    """
    Create a Word document from parsed sessions.
    
    Args:
        sessions: List of parsed session dictionaries
        
    Returns:
        python-docx Document object
    """
    doc = Document()
    
    # Set document title
    title = doc.add_heading('NeuralDBG - Session Summaries', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation timestamp
    timestamp = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # Spacer
    
    # Process each session (newest first is already in the file)
    for session in sessions:
        # Session header
        doc.add_heading(f'Session Summary — {session["date"]}', level=1)
        
        # Editor
        if session['editor']:
            editor_para = doc.add_paragraph()
            editor_run = editor_para.add_run(f'Editor: {session["editor"]}')
            editor_run.bold = True
        
        # French section
        if session['francais']:
            add_language_section(doc, 'Francais', session['francais'])
        
        # English section
        if session['english']:
            add_language_section(doc, 'English', session['english'])
        
        # Metadata (Tests, Blockers, Progress)
        meta_para = doc.add_paragraph()
        
        if session['tests']:
            tests_run = meta_para.add_run(f'Tests: {session["tests"]}   ')
            tests_run.bold = True
        
        if session['blockers']:
            blockers_run = meta_para.add_run(f'Blockers: {session["blockers"]}   ')
            blockers_run.bold = True
        
        if session['progress']:
            progress_run = meta_para.add_run(f'Progress: {session["progress"]}')
            progress_run.bold = True
            # Color based on progress
            try:
                progress_val = int(session['progress'].replace('%', ''))
                if progress_val >= 80:
                    progress_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                elif progress_val >= 50:
                    progress_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                else:
                    progress_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            except ValueError:
                pass
        
        # Separator
        doc.add_paragraph('─' * 40)
    
    return doc


def add_language_section(doc: Document, lang_name: str, section_data: dict) -> None:
    """Add a language section to the document."""
    doc.add_heading(lang_name, level=2)
    
    field_names = {
        'what_done': 'What was done' if lang_name == 'English' else 'Ce qui a ete fait',
        'initiatives': 'Initiatives given' if lang_name == 'English' else 'Initiatives donnees',
        'files_changed': 'Files changed' if lang_name == 'English' else 'Fichiers modifies',
        'next_steps': 'Next steps' if lang_name == 'English' else 'Etapes suivantes'
    }
    
    for key, items in section_data.items():
        if items:
            para = doc.add_paragraph()
            para.add_run(f'{field_names[key]}:').bold = True
            
            for item in items:
                doc.add_paragraph(item, style='List Bullet')


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description='Convert SESSION_SUMMARY.md to Word document (.docx)'
    )
    parser.add_argument(
        '--input', '-i',
        default='SESSION_SUMMARY.md',
        help='Input markdown file (default: SESSION_SUMMARY.md)'
    )
    parser.add_argument(
        '--output', '-o',
        default=None,
        help='Output docx file (default: SESSION_SUMMARY_YYYY-MM-DD.docx)'
    )
    parser.add_argument(
        '--latest', '-l',
        action='store_true',
        help='Only include the latest session'
    )
    
    args = parser.parse_args()
    
    # Read input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: Input file not found: {input_path}")
        sys.exit(1)
    
    markdown_content = input_path.read_text(encoding='utf-8')
    
    # Parse sessions
    sessions = parse_session_summary(markdown_content)
    
    if not sessions:
        print("WARNING: No sessions found in input file.")
        sys.exit(0)
    
    # Filter to latest if requested
    if args.latest:
        sessions = [sessions[0]]
    
    # Create document
    doc = create_document(sessions)
    
    # Determine output path
    if args.output:
        output_path = Path(args.output)
    else:
        date_str = datetime.now().strftime('%Y-%m-%d')
        output_path = Path(f'SESSION_SUMMARY_{date_str}.docx')
    
    # Save document
    doc.save(output_path)
    print(f"SUCCESS: Created {output_path}")
    print(f"  - {len(sessions)} session(s) exported")


if __name__ == '__main__':
    main()